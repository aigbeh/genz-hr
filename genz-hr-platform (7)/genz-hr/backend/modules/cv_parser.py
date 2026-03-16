"""
GENZ HR — CV Parser & Candidate Scorer
Extracts structured data from PDF/DOCX CVs and scores candidates.
"""
import re
import json
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

import pdfplumber
from docx import Document


EDUCATION_KEYWORDS = {
    "phd": 10, "doctorate": 10,
    "msc": 8, "mba": 8, "masters": 8, "m.sc": 8, "m.eng": 8,
    "bsc": 6, "b.sc": 6, "hnd": 5, "beng": 6, "b.eng": 6,
    "ond": 3, "diploma": 3, "certificate": 2,
    "first class": 3, "second class upper": 2, "second class lower": 1,
}

EXPERIENCE_PATTERNS = [
    r"(\d+)\+?\s*years?\s*(?:of\s*)?experience",
    r"(\d+)\+?\s*yrs?\s*(?:of\s*)?experience",
    r"experience\s*of\s*(\d+)\+?\s*years?",
]

SECTION_HEADERS = [
    "experience", "work experience", "employment history",
    "education", "qualifications", "skills", "technical skills",
    "certifications", "summary", "objective", "projects",
]


@dataclass
class ParsedCV:
    name: str = ""
    email: str = ""
    phone: str = ""
    years_experience: int = 0
    education_level: str = ""
    skills: list[str] = field(default_factory=list)
    companies_worked: list[str] = field(default_factory=list)
    raw_text: str = ""
    sections: dict = field(default_factory=dict)


@dataclass  
class ScoredCandidate:
    name: str
    email: str
    phone: str
    position_applied: str
    cv_path: str
    education_score: float  # 0–100
    skills_score: float
    experience_score: float
    keyword_score: float
    total_score: float
    raw_cv_text: str
    ai_summary: str = ""
    
    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "position_applied": self.position_applied,
            "cv_path": self.cv_path,
            "education_score": round(self.education_score, 1),
            "skills_score": round(self.skills_score, 1),
            "experience_score": round(self.experience_score, 1),
            "keyword_score": round(self.keyword_score, 1),
            "total_score": round(self.total_score, 1),
            "ai_summary": self.ai_summary,
        }


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        text = f"[Error reading PDF: {e}]"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"[Error reading DOCX: {e}]"
    return text


def extract_text(file_path: str) -> str:
    """Auto-detect format and extract text."""
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return extract_text_from_pdf(file_path)
    elif path.suffix.lower() in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def parse_cv(text: str) -> ParsedCV:
    """Extract structured data from raw CV text."""
    cv = ParsedCV(raw_text=text)
    text_lower = text.lower()

    # Email
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    if email_match:
        cv.email = email_match.group()

    # Phone (Nigerian patterns)
    phone_match = re.search(
        r"(?:\+234|0)[\s-]?[789]\d{1}[\s-]?\d{4}[\s-]?\d{4}", text
    )
    if phone_match:
        cv.phone = phone_match.group().strip()

    # Years experience
    for pattern in EXPERIENCE_PATTERNS:
        match = re.search(pattern, text_lower)
        if match:
            cv.years_experience = int(match.group(1))
            break

    # Education level
    for keyword, _ in sorted(EDUCATION_KEYWORDS.items(), key=lambda x: -x[1]):
        if keyword in text_lower:
            cv.education_level = keyword.upper()
            break

    # Skills (simple keyword extraction from skills section)
    skills_section = _extract_section(text, ["skills", "technical skills", "competencies"])
    if skills_section:
        # Extract comma/newline separated items
        raw_skills = re.split(r"[,\n•·|]", skills_section)
        cv.skills = [s.strip() for s in raw_skills if 2 < len(s.strip()) < 50][:20]

    return cv


def _extract_section(text: str, headers: list[str]) -> str:
    """Extract text under a named section."""
    text_lower = text.lower()
    for header in headers:
        pattern = rf"\b{re.escape(header)}\b[:\s]*\n(.*?)(?:\n[A-Z]{{3,}}|\Z)"
        match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
        if match:
            return match.group(1)[:1000]
    return ""


def score_education(cv: ParsedCV) -> float:
    """Score 0–100 based on highest education level."""
    edu_lower = cv.education_level.lower()
    for keyword, points in EDUCATION_KEYWORDS.items():
        if keyword in edu_lower:
            return min(100, points * 10)
    # Fallback: look in raw text
    for keyword, points in sorted(EDUCATION_KEYWORDS.items(), key=lambda x: -x[1]):
        if keyword in cv.raw_text.lower():
            return min(100, points * 10)
    return 20  # default if we find text but no match


def score_experience(cv: ParsedCV, required_years: int = 3) -> float:
    """Score 0–100 based on years of experience vs required."""
    if required_years <= 0:
        return 50
    ratio = cv.years_experience / required_years
    return min(100, ratio * 70)  # max 70 from years alone, leaves room for quality


def score_skills(cv: ParsedCV, required_skills: list[str]) -> float:
    """Score 0–100 based on matched required skills."""
    if not required_skills:
        return 50
    cv_text_lower = cv.raw_text.lower()
    matched = sum(1 for skill in required_skills if skill.lower() in cv_text_lower)
    return min(100, (matched / len(required_skills)) * 100)


def score_keywords(cv: ParsedCV, role_keywords: list[str]) -> float:
    """Score 0–100 based on role-specific keyword density."""
    if not role_keywords:
        return 50
    cv_text_lower = cv.raw_text.lower()
    matched = sum(1 for kw in role_keywords if kw.lower() in cv_text_lower)
    return min(100, (matched / len(role_keywords)) * 100)


def score_candidate(
    file_path: str,
    position: str,
    required_skills: list[str] = None,
    role_keywords: list[str] = None,
    required_years: int = 3,
    weights: dict = None,
) -> ScoredCandidate:
    """
    Full CV scoring pipeline.
    
    Weights default to:
        education: 25%
        skills: 35%
        experience: 25%
        keywords: 15%
    """
    if weights is None:
        weights = {
            "education": 0.25,
            "skills": 0.35,
            "experience": 0.25,
            "keywords": 0.15,
        }

    required_skills = required_skills or []
    role_keywords = role_keywords or []

    text = extract_text(file_path)
    cv = parse_cv(text)

    edu_score = score_education(cv)
    exp_score = score_experience(cv, required_years)
    skills_score = score_skills(cv, required_skills)
    kw_score = score_keywords(cv, role_keywords)

    total = (
        edu_score * weights["education"]
        + skills_score * weights["skills"]
        + exp_score * weights["experience"]
        + kw_score * weights["keywords"]
    )

    return ScoredCandidate(
        name=cv.name or Path(file_path).stem,
        email=cv.email,
        phone=cv.phone,
        position_applied=position,
        cv_path=file_path,
        education_score=edu_score,
        skills_score=skills_score,
        experience_score=exp_score,
        keyword_score=kw_score,
        total_score=round(total, 2),
        raw_cv_text=text[:5000],
    )


def rank_candidates(candidates: list[ScoredCandidate], top_n: int = 5) -> list[ScoredCandidate]:
    """Sort candidates by total score and assign ranks."""
    sorted_candidates = sorted(candidates, key=lambda c: c.total_score, reverse=True)
    for i, c in enumerate(sorted_candidates):
        # Mutate rank in-place (add rank attribute)
        c.__dict__["rank"] = i + 1
    return sorted_candidates


def generate_shortlist_report(candidates: list[ScoredCandidate], position: str) -> str:
    """Generate a markdown summary of the shortlist."""
    lines = [
        f"# GENZ HR — Recruitment Shortlist",
        f"**Position:** {position}",
        f"**Total Applicants:** {len(candidates)}",
        f"**Shortlisted:** {sum(1 for c in candidates if c.total_score >= 60)}",
        "",
        "| Rank | Name | Email | Total Score | Education | Skills | Experience |",
        "|------|------|-------|-------------|-----------|--------|------------|",
    ]
    for c in candidates[:10]:
        rank = c.__dict__.get("rank", "—")
        shortlist_marker = "✅" if c.total_score >= 60 else "❌"
        lines.append(
            f"| {rank} {shortlist_marker} | {c.name} | {c.email} | "
            f"**{c.total_score:.1f}** | {c.education_score:.0f} | "
            f"{c.skills_score:.0f} | {c.experience_score:.0f} |"
        )

    lines += [
        "",
        "> _Scores are AI-generated. Esther may override any ranking before shortlist is finalized._",
        "",
        "## AI Recommendation",
        "Candidates scoring 60+ are recommended for interview. "
        "Please review profiles and adjust scores if needed before proceeding.",
    ]
    return "\n".join(lines)
